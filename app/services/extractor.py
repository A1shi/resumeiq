import os
import pypdf
import docx

LIGATURE_MAP = {
    '\ufb00': 'ff',
    '\ufb01': 'fi',
    '\ufb02': 'fl',
    '\ufb03': 'ffi',
    '\ufb04': 'ffl',
    '\ufb05': 'ft',
    '\ufb06': 'st'
}

def clean_ligatures(text: str) -> str:
    if not text:
        return text
    for ligature, replacement in LIGATURE_MAP.items():
        text = text.replace(ligature, replacement)
    return text

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract raw text content from a PDF file using pypdf.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    text_content = []
    try:
        reader = pypdf.PdfReader(file_path)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            page_text = page.extract_text()
            if page_text:
                text_content.append(clean_ligatures(page_text))
    except Exception as e:
        raise ValueError(f"Failed to parse PDF document: {str(e)}")
        
    return "\n\n".join(text_content).strip()

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract raw text content from a DOCX Word file using python-docx.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    try:
        doc = docx.Document(file_path)
        text_content = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text:
                text_content.append(para.text)
                
        # Extract from tables (resumes often use tables for formatting)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
                    
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX document: {str(e)}")
        
    return "\n".join(text_content).strip()

def extract_text_from_image_gemini(file_path: str) -> str:
    """
    Extract raw text from a JPG/JPEG image using Gemini Multimodal capability.
    """
    from PIL import Image
    import google.generativeai as genai
    from app.config import settings

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        img = Image.open(file_path)
        genai.configure(api_key=settings.GEMINI_API_KEY)
        model = genai.GenerativeModel("gemini-1.5-flash")
        
        response = model.generate_content([
            "Extract all readable text from this resume image. Output only the plain text found in the image. Keep the original structure as much as possible.",
            img
        ])
        if response.text:
            return response.text.strip()
        raise ValueError("Gemini returned empty text for the image.")
    except Exception as e:
        # Fallback error message if Gemini OCR fails/is offline
        raise ValueError(f"Failed to perform OCR on image: {str(e)}")

def extract_text(file_path: str) -> str:
    """
    General purpose text extraction router based on file extension.
    """
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif ext in [".jpg", ".jpeg", ".png"]:
        return extract_text_from_image_gemini(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, and JPG/JPEG are supported.")
