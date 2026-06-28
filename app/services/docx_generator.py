import io
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from typing import Dict, Any, Optional

LIGATURE_MAP = {
    '\ufb00': 'ff',
    '\ufb01': 'fi',
    '\ufb02': 'fl',
    '\ufb03': 'ffi',
    '\ufb04': 'ffl',
    '\ufb05': 'ft',
    '\ufb06': 'st'
}

def clean_ligatures(text: Any) -> Any:
    if isinstance(text, str):
        for ligature, replacement in LIGATURE_MAP.items():
            text = text.replace(ligature, replacement)
        return text
    elif isinstance(text, dict):
        return {k: clean_ligatures(v) for k, v in text.items()}
    elif isinstance(text, list):
        return [clean_ligatures(item) for item in text]
    return text

def generate_cover_letter_docx(cover_letter_text: str) -> io.BytesIO:
    """
    Generates a professional DOCX cover letter from plain text.
    Returns a bytes buffer containing the DOCX file.
    """
    cover_letter_text = clean_ligatures(cover_letter_text)
    doc = Document()
    
    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Standard styles setup
    style = doc.styles['Normal']
    font = style.font  # type: ignore
    font.name = 'Arial'
    font.size = Pt(11)
    
    paragraphs = cover_letter_text.split("\n\n")
    
    # Clean and split
    lines = [line.strip() for line in cover_letter_text.splitlines() if line.strip()]
    
    if len(lines) >= 2 and "|" in lines[1]:
        # Professional header layout
        name_p = doc.add_paragraph()
        name_p.paragraph_format.space_after = Pt(2)
        name_run = name_p.add_run(lines[0])
        name_run.bold = True
        name_run.font.size = Pt(14)
        
        contact_p = doc.add_paragraph()
        contact_p.paragraph_format.space_after = Pt(24) # Space after header block
        contact_run = contact_p.add_run(lines[1])
        contact_run.font.size = Pt(10)
        
        # Add the rest of the paragraphs
        body_text = "\n\n".join(paragraphs[1:])
        for p_text in body_text.split("\n\n"):
            p_text = p_text.strip()
            if p_text:
                p = doc.add_paragraph(p_text)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = 1.15
    else:
        for p_text in paragraphs:
            p_text = p_text.strip()
            if p_text:
                p = doc.add_paragraph(p_text)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = 1.15
                
    # Save to a bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def categorize_skills(skills_list):
    categories = {
        "Languages": [],
        "Databases & Tools": [],
        "Concepts & Methods": []
    }
    
    languages_keywords = {"python", "sql", "r", "julia", "matlab", "sas", "scala", "c++", "java", "javascript", "typescript", "bash"}
    databases_tools = {"postgresql", "mysql", "mongodb", "snowflake", "redshift", "excel", "tableau", "power bi", "powerbi", "spark", "hadoop", "oracle", "bigquery", "aws", "gcp", "azure", "docker", "kubernetes", "git", "pandas", "numpy", "matplotlib", "seaborn", "scikit-learn", "keras", "tensorflow", "pytorch", "jupyter"}
    
    for s in skills_list:
        s_lower = str(s).lower().strip()
        if any(kw in s_lower for kw in languages_keywords):
            categories["Languages"].append(s)
        elif any(kw in s_lower for kw in databases_tools):
            categories["Databases & Tools"].append(s)
        else:
            categories["Concepts & Methods"].append(s)
            
    # Remove empty categories and return
    return {k: v for k, v in categories.items() if v}


def generate_resume_template_docx(resume_data: Dict[str, Any], template_name: str, customization: Optional[Dict[str, Any]] = None) -> io.BytesIO:
    """
    Generates a beautifully formatted resume DOCX strictly based on the selected
    template style and custom candidate data.
    """
    resume_data = clean_ligatures(resume_data)
    doc = Document()
    
    if customization is None:
        customization = resume_data.get("customization") or {}
        
    # Page setup - margins customization (default 0.5 inch)
    margin_val = 0.5
    cust_margin = customization.get("marginSize")
    if cust_margin is not None:
        try:
            margin_val = float(cust_margin) / 72.0
        except (ValueError, TypeError):
            pass
            
    for section in doc.sections:
        section.top_margin = Inches(margin_val)
        section.bottom_margin = Inches(margin_val)
        section.left_margin = Inches(margin_val)
        section.right_margin = Inches(margin_val)
        
    # Helpers for margins, borders and shading
    def set_cell_shading(cell, color_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def add_bottom_border(paragraph, color_hex="CBD5E1"):
        pPr = paragraph._p.get_or_add_pPr()
        pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="{color_hex}"/></w:pBdr>')
        pPr.append(pbdr)

    def safe_get(obj, key, default=None):
        if not obj:
            return default
        if isinstance(obj, dict):
            return obj.get(key, default)
        return getattr(obj, key, default)

    # Set default fonts & colors based on template
    font_name = 'Arial'
    primary_color = RGBColor(0x1e, 0x29, 0x3b) # slate-800
    accent_color = RGBColor(0x4f, 0x46, 0xe5)  # indigo
    border_color_hex = "CBD5E1"

    if template_name == "ATS Professional":
        font_name = 'Arial'
        primary_color = RGBColor(0x0f, 0x17, 0x2a)
        accent_color = RGBColor(0x33, 0x41, 0x55)
        border_color_hex = "CBD5E1"
    elif template_name == "Modern Professional":
        font_name = 'Calibri'
        primary_color = RGBColor(0x0f, 0x17, 0x2a)
        accent_color = RGBColor(0x25, 0x63, 0xeb)
        border_color_hex = "93C5FD"
    elif template_name in ["Creative", "Creative Designer", "Creative Resume"]:
        font_name = 'Arial'
        primary_color = RGBColor(0x1e, 0x1b, 0x4b)
        accent_color = RGBColor(0xec, 0x48, 0x99)
        border_color_hex = "F472B6"
    elif template_name == "Software Engineer":
        font_name = 'Segoe UI'
        primary_color = RGBColor(0x0f, 0x17, 0x2a)
        accent_color = RGBColor(0x63, 0x66, 0xf1)
        border_color_hex = "CBD5E1"
    elif template_name == "Data Analyst":
        font_name = 'Calibri'
        primary_color = RGBColor(0x1f, 0x29, 0x37)
        accent_color = RGBColor(0x0f, 0x76, 0x6e)
        border_color_hex = "99F6E4"
    elif template_name in ["Executive", "Executive Resume"]:
        font_name = 'Times New Roman'
        primary_color = RGBColor(0x0f, 0x17, 0x2a)
        accent_color = RGBColor(0x1e, 0x3a, 0x8a)
        border_color_hex = "D97706" # Gold
    elif template_name == "Minimal Elegant":
        font_name = 'Georgia'
        primary_color = RGBColor(0x09, 0x09, 0x0b)
        accent_color = RGBColor(0x71, 0x71, 0x7a)
        border_color_hex = "E4E4E7"
    elif template_name in ["Student/Fresher", "Student / Fresher"]:
        font_name = 'Georgia'
        primary_color = RGBColor(0x1e, 0x29, 0x3b)
        accent_color = RGBColor(0x10, 0xb9, 0x81)
        border_color_hex = "A7F3D0"

    # Extract fields
    name = safe_get(resume_data, "name") or "Candidate Name"
    email = safe_get(resume_data, "email") or "email@example.com"
    phone = safe_get(resume_data, "phone") or "Phone Not Provided"
    skills = safe_get(resume_data, "skills") or []
    experience = safe_get(resume_data, "experience") or []
    projects = safe_get(resume_data, "projects") or []
    education = safe_get(resume_data, "education") or []
    certifications = safe_get(resume_data, "certifications") or []
    languages = safe_get(resume_data, "languages") or []
    summary_text = safe_get(resume_data, "professional_summary") or ""

    # Dual column layouts (Modern Professional & Creative Designer / Creative Resume)
    if template_name in ["Modern Professional", "Creative", "Creative Designer", "Creative Resume"]:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        
        # Modern: sidebar left (2.2 in), content right (5.3 in). Creative: content left (5.0 in), sidebar right (2.5 in).
        if template_name == "Modern Professional":
            table.columns[0].width = Inches(2.2)
            table.columns[1].width = Inches(5.3)
            left_cell = table.cell(0, 0)
            right_cell = table.cell(0, 1)
            
            set_cell_shading(left_cell, "F8FAFC")
            set_cell_margins(left_cell, top=144, bottom=144, left=150, right=150)
            set_cell_margins(right_cell, top=144, bottom=144, left=200, right=150)
        else:
            table.columns[0].width = Inches(5.0)
            table.columns[1].width = Inches(2.5)
            left_cell = table.cell(0, 1) # sidebar column
            right_cell = table.cell(0, 0) # content column
            
            set_cell_shading(left_cell, "FFF1F2") # soft pink shading
            set_cell_margins(left_cell, top=144, bottom=144, left=150, right=150)
            set_cell_margins(right_cell, top=144, bottom=144, left=150, right=200)
        
        # Helper to add paragraph in a cell
        def add_cell_p(cell, text="", bold=False, size=10, color=primary_color, space_before=0, space_after=2, italic=False):
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.line_spacing = 1.15
            if text:
                run = p.add_run(text)
                run.bold = bold
                run.italic = italic
                run.font.name = font_name
                run.font.size = Pt(size)
                run.font.color.rgb = color
            return p
            
        # Add sidebar content (left_cell)
        add_cell_p(left_cell, "CONTACT INFO", bold=True, size=10.5, color=accent_color, space_before=6, space_after=4)
        add_cell_p(left_cell, f"✉ {email}", size=9, space_after=3)
        add_cell_p(left_cell, f"☎ {phone}", size=9, space_after=12)
        
        if skills:
            add_cell_p(left_cell, "SKILLS", bold=True, size=10.5, color=accent_color, space_before=10, space_after=4)
            if template_name == "Modern Professional":
                for idx, s in enumerate(skills):
                    percentages = [90, 85, 80, 75, 95]
                    pct = percentages[idx % len(percentages)]
                    bar_len = int(pct / 10)
                    bar_str = "█" * bar_len + "░" * (10 - bar_len)
                    add_cell_p(left_cell, f"• {s} [{bar_str}] {pct}%", size=8.5, space_after=3)
            else:
                for s in skills:
                    add_cell_p(left_cell, f"• {s}", size=9, space_after=3)
                
        if languages:
            add_cell_p(left_cell, "LANGUAGES", bold=True, size=10.5, color=accent_color, space_before=10, space_after=4)
            for lang in languages:
                l_name = safe_get(lang, "language")
                prof = safe_get(lang, "proficiency")
                prof_str = f" ({prof})" if prof else ""
                add_cell_p(left_cell, f"• {l_name}{prof_str}", size=9, space_after=3)
                
        if certifications:
            add_cell_p(left_cell, "CERTIFICATIONS", bold=True, size=10.5, color=accent_color, space_before=10, space_after=4)
            for cert in certifications:
                c_name = safe_get(cert, "name")
                add_cell_p(left_cell, f"• {c_name}", size=9, space_after=3)
                
        # Add main column content (right_cell)
        if template_name == "Modern Professional":
            # Profile Card in a shaded single-cell nested table
            card_table = right_cell.add_table(rows=1, cols=1)
            card_table.autofit = False
            card_table.columns[0].width = Inches(5.1)
            card_cell = card_table.cell(0, 0)
            set_cell_shading(card_cell, "F8FAFC")
            set_cell_margins(card_cell, top=100, bottom=100, left=150, right=150)
            
            p_card_name = card_cell.add_paragraph()
            p_card_name.paragraph_format.space_before = Pt(4)
            p_card_name.paragraph_format.space_after = Pt(2)
            r_card_name = p_card_name.add_run(name.upper())
            r_card_name.bold = True
            r_card_name.font.name = font_name
            r_card_name.font.size = Pt(14)
            r_card_name.font.color.rgb = primary_color
            
            p_card_sub = card_cell.add_paragraph()
            p_card_sub.paragraph_format.space_after = Pt(4)
            r_card_sub = p_card_sub.add_run("PROFESSIONAL PROFILE")
            r_card_sub.bold = True
            r_card_sub.font.name = font_name
            r_card_sub.font.size = Pt(8.5)
            r_card_sub.font.color.rgb = accent_color
            
            if summary_text:
                p_card_sum = card_cell.add_paragraph()
                p_card_sum.paragraph_format.space_after = Pt(4)
                r_card_sum = p_card_sum.add_run(summary_text)
                r_card_sum.font.name = font_name
                r_card_sum.font.size = Pt(9)
                r_card_sum.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        else:
            add_cell_p(right_cell, name, bold=True, size=20, color=primary_color, space_after=2)
            if template_name in ["Creative", "Creative Designer", "Creative Resume"]:
                add_cell_p(right_cell, "CREATIVE PROFESSIONAL", bold=True, size=9.5, color=accent_color, space_after=8)
            
            if summary_text:
                p_sec = add_cell_p(right_cell, "PROFESSIONAL SUMMARY", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
                add_bottom_border(p_sec, border_color_hex)
                add_cell_p(right_cell, summary_text, size=9.5, space_after=8)

        # Experience
        if experience:
            p_sec = add_cell_p(right_cell, "WORK EXPERIENCE", bold=True, size=12, color=accent_color, space_before=12, space_after=4)
            add_bottom_border(p_sec, border_color_hex)
            
            for exp in experience:
                role = safe_get(exp, "role") or "Role"
                company = safe_get(exp, "company") or "Company"
                start = safe_get(exp, "start_date") or "N/A"
                end = safe_get(exp, "end_date") or "Present"
                desc = safe_get(exp, "description") or ""
                
                p_job = right_cell.add_paragraph()
                p_job.paragraph_format.space_before = Pt(4)
                p_job.paragraph_format.space_after = Pt(2)
                p_job.paragraph_format.line_spacing = 1.15
                
                r_title = p_job.add_run(f"{role}")
                r_title.bold = True
                r_title.font.name = font_name
                r_title.font.size = Pt(9.5)
                
                r_comp = p_job.add_run(f" at {company} ({start} - {end})")
                r_comp.font.name = font_name
                r_comp.font.size = Pt(9)
                r_comp.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                
                if desc:
                    for bullet in desc.split("\n"):
                        if bullet.strip():
                            b_text = bullet.strip().lstrip("-").lstrip("•").strip()
                            add_cell_p(right_cell, f"• {b_text}", size=8.5, space_after=2)
                            
        # Projects
        if projects:
            p_sec = add_cell_p(right_cell, "PROJECTS", bold=True, size=12, color=accent_color, space_before=12, space_after=4)
            add_bottom_border(p_sec, border_color_hex)
            
            for proj in projects:
                title = safe_get(proj, "title") or "Project Title"
                desc = safe_get(proj, "description") or ""
                tech_list = safe_get(proj, "technologies") or []
                tech = ", ".join(tech_list)
                tech_suffix = f" ({tech})" if tech else ""
                
                p_proj = right_cell.add_paragraph()
                p_proj.paragraph_format.space_before = Pt(4)
                p_proj.paragraph_format.space_after = Pt(2)
                
                r_title = p_proj.add_run(f"{title}")
                r_title.bold = True
                r_title.font.name = font_name
                r_title.font.size = Pt(9.5)
                
                if tech_suffix:
                    r_tech = p_proj.add_run(tech_suffix)
                    r_tech.font.name = font_name
                    r_tech.font.size = Pt(8.5)
                    r_tech.italic = True
                    r_tech.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                    
                if desc:
                    for bullet in desc.split("\n"):
                        if bullet.strip():
                            b_text = bullet.strip().lstrip("-").lstrip("•").strip()
                            add_cell_p(right_cell, f"• {b_text}", size=8.5, space_after=2)
                            
        # Education
        if education:
            p_sec = add_cell_p(right_cell, "EDUCATION", bold=True, size=12, color=accent_color, space_before=12, space_after=4)
            add_bottom_border(p_sec, border_color_hex)
            
            for edu in education:
                degree = safe_get(edu, "degree") or "Degree"
                major = safe_get(edu, "field_of_study") or ""
                school = safe_get(edu, "school") or "School"
                end = safe_get(edu, "end_date") or ""
                
                degree_str = f"{degree} in {major}" if major else degree
                
                p_edu = right_cell.add_paragraph()
                p_edu.paragraph_format.space_before = Pt(4)
                p_edu.paragraph_format.space_after = Pt(2)
                
                r_deg = p_edu.add_run(f"{degree_str}")
                r_deg.bold = True
                r_deg.font.name = font_name
                r_deg.font.size = Pt(9.5)
                
                r_sch = p_edu.add_run(f"  —  {school} (Graduated: {end})" if end else f"  —  {school}")
                r_sch.font.name = font_name
                r_sch.font.size = Pt(9)
                r_sch.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    else:
        # Single column layout for other templates
        header_align = WD_ALIGN_PARAGRAPH.LEFT
        if template_name in ["Executive", "Executive Resume", "Student/Fresher", "Student / Fresher"]:
            header_align = WD_ALIGN_PARAGRAPH.CENTER
            
        p_name = doc.add_paragraph()
        p_name.alignment = header_align
        p_name.paragraph_format.space_after = Pt(2)
        run_name = p_name.add_run(name)
        run_name.bold = True
        run_name.font.name = font_name
        run_name.font.size = Pt(22)
        run_name.font.color.rgb = primary_color
        
        p_contact = doc.add_paragraph()
        p_contact.alignment = header_align
        p_contact.paragraph_format.space_after = Pt(12)
        
        if template_name == "Software Engineer":
            github_slug = name.lower().replace(" ", "").replace("-", "")
            contact_text = f"// {email}   |   {phone}   |   github.com/{github_slug}"
        else:
            parts = [email, phone]
            contact_text = "  |  ".join(parts)
            
        run_contact = p_contact.add_run(contact_text)
        run_contact.font.name = font_name
        run_contact.font.size = Pt(9.5)
        run_contact.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        
        def add_p(text="", bold=False, size=9.5, color=primary_color, space_before=0, space_after=3, italic=False):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.line_spacing = 1.15
            if text:
                run = p.add_run(text)
                run.bold = bold
                run.italic = italic
                run.font.name = font_name
                run.font.size = Pt(size)
                run.font.color.rgb = color
            return p

        # Summary
        if summary_text:
            p_sec = add_p("PROFESSIONAL SUMMARY", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
            add_bottom_border(p_sec, border_color_hex)
            add_p(summary_text, size=9.5, space_after=8)
            
        # Determine section order based on template/customization
        sections = customization.get("section_order") or resume_data.get("section_order")
        if not sections:
            sections = ["experience", "projects", "skills", "education", "cert_lang"]
            if template_name in ["Student/Fresher", "Student / Fresher"]:
                sections = ["education", "projects", "skills", "experience", "cert_lang"]
            elif template_name == "Data Analyst":
                sections = ["skills", "experience", "projects", "education", "cert_lang"]
            elif template_name == "Software Engineer":
                sections = ["skills", "projects", "experience", "education", "cert_lang"]
                
        for sec in sections:
            if sec == "experience" and experience:
                p_sec = add_p("WORK EXPERIENCE", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
                add_bottom_border(p_sec, border_color_hex)
                
                for exp in experience:
                    role = safe_get(exp, "role") or "Role"
                    company = safe_get(exp, "company") or "Company"
                    start = safe_get(exp, "start_date") or "N/A"
                    end = safe_get(exp, "end_date") or "Present"
                    desc = safe_get(exp, "description") or ""
                    
                    p_job = doc.add_paragraph()
                    p_job.paragraph_format.space_before = Pt(4)
                    p_job.paragraph_format.space_after = Pt(2)
                    p_job.paragraph_format.line_spacing = 1.15
                    
                    r_title = p_job.add_run(f"{role}")
                    r_title.bold = True
                    r_title.font.name = font_name
                    r_title.font.size = Pt(10)
                    
                    r_comp = p_job.add_run(f"  —  {company} ({start} - {end})")
                    r_comp.font.name = font_name
                    r_comp.font.size = Pt(9.5)
                    r_comp.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                    
                    if desc:
                        for bullet in desc.split("\n"):
                            if bullet.strip():
                                b_text = bullet.strip().lstrip("-").lstrip("•").strip()
                                add_p(f"• {b_text}", size=9, space_after=2)
                                
            elif sec == "projects" and projects:
                p_sec = add_p("PROJECTS", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
                add_bottom_border(p_sec, border_color_hex)
                
                for proj in projects:
                    title = safe_get(proj, "title") or "Project Title"
                    desc = safe_get(proj, "description") or ""
                    tech_list = safe_get(proj, "technologies") or []
                    tech = ", ".join(tech_list)
                    tech_suffix = f" ({tech})" if tech else ""
                    
                    p_proj = doc.add_paragraph()
                    p_proj.paragraph_format.space_before = Pt(4)
                    p_proj.paragraph_format.space_after = Pt(2)
                    
                    r_title = p_proj.add_run(f"{title}")
                    r_title.bold = True
                    r_title.font.name = font_name
                    r_title.font.size = Pt(10)
                    
                    if tech_suffix:
                        r_tech = p_proj.add_run(tech_suffix)
                        r_tech.font.name = font_name
                        r_tech.font.size = Pt(9)
                        r_tech.italic = True
                        r_tech.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                        
                    if desc:
                        for bullet in desc.split("\n"):
                            if bullet.strip():
                                b_text = bullet.strip().lstrip("-").lstrip("•").strip()
                                add_p(f"• {b_text}", size=9, space_after=2)
                                
            elif sec == "skills" and skills:
                if template_name == "Data Analyst":
                    p_sec = add_p("TECHNICAL SKILLS MATRIX", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
                    add_bottom_border(p_sec, border_color_hex)
                    cats = categorize_skills(skills)
                    for cat_name, cat_skills in cats.items():
                        skills_str = ", ".join(cat_skills)
                        p_cat = doc.add_paragraph()
                        p_cat.paragraph_format.space_before = Pt(2)
                        p_cat.paragraph_format.space_after = Pt(2)
                        r_lbl = p_cat.add_run(f"{cat_name}: ")
                        r_lbl.bold = True
                        r_lbl.font.name = font_name
                        r_lbl.font.size = Pt(9.5)
                        r_lbl.font.color.rgb = accent_color
                        
                        r_val = p_cat.add_run(skills_str)
                        r_val.font.name = font_name
                        r_val.font.size = Pt(9.5)
                        r_val.font.color.rgb = primary_color
                else:
                    p_sec = add_p("TECHNICAL SKILLS", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
                    add_bottom_border(p_sec, border_color_hex)
                    
                    if template_name == "Software Engineer":
                        p_sk = doc.add_paragraph()
                        p_sk.paragraph_format.space_after = Pt(6)
                        for s in skills:
                            r_s = p_sk.add_run(f"{s}, ")
                            r_s.font.name = 'Courier New'
                            r_s.font.size = Pt(9.5)
                            r_s.font.color.rgb = primary_color
                    else:
                        skills_text = ", ".join(skills)
                        add_p(skills_text, size=9.5, space_after=6)
                
            elif sec == "education" and education:
                p_sec = add_p("EDUCATION", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
                add_bottom_border(p_sec, border_color_hex)
                
                for edu in education:
                    degree = safe_get(edu, "degree") or "Degree"
                    major = safe_get(edu, "field_of_study") or ""
                    school = safe_get(edu, "school") or "School"
                    end = safe_get(edu, "end_date") or ""
                    
                    degree_str = f"{degree} in {major}" if major else degree
                    
                    p_edu = doc.add_paragraph()
                    p_edu.paragraph_format.space_before = Pt(4)
                    p_edu.paragraph_format.space_after = Pt(2)
                    
                    r_deg = p_edu.add_run(f"{degree_str}")
                    r_deg.bold = True
                    r_deg.font.name = font_name
                    r_deg.font.size = Pt(10)
                    
                    r_sch = p_edu.add_run(f"  —  {school} ({end})" if end else f"  —  {school}")
                    r_sch.font.name = font_name
                    r_sch.font.size = Pt(9.5)
                    r_sch.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                    
            elif sec == "certifications" and certifications:
                p_sec = add_p("CERTIFICATIONS", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
                add_bottom_border(p_sec, border_color_hex)
                
                cert_items = []
                for cert in certifications:
                    c_name = safe_get(cert, "name")
                    c_issuer = safe_get(cert, "issuer")
                    issuer_str = f" ({c_issuer})" if c_issuer else ""
                    cert_items.append(f"{c_name}{issuer_str}")
                if cert_items:
                    add_p(", ".join(cert_items), size=9.5, space_after=6)
                    
            elif sec == "languages" and languages:
                p_sec = add_p("LANGUAGES", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
                add_bottom_border(p_sec, border_color_hex)
                
                lang_items = []
                for lang in languages:
                    l_name = safe_get(lang, "language")
                    prof = safe_get(lang, "proficiency")
                    prof_str = f" ({prof})" if prof else ""
                    lang_items.append(f"{l_name}{prof_str}")
                if lang_items:
                    add_p(", ".join(lang_items), size=9.5, space_after=6)
                    
            elif sec == "achievements" and resume_data.get("achievements"):
                p_sec = add_p("ACHIEVEMENTS", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
                add_bottom_border(p_sec, border_color_hex)
                
                for ach in resume_data["achievements"]:
                    add_p(f"• {ach}", size=9, space_after=2)
                    
            elif sec == "interests" and resume_data.get("interests"):
                p_sec = add_p("INTERESTS & HOBBIES", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
                add_bottom_border(p_sec, border_color_hex)
                add_p(", ".join(resume_data["interests"]), size=9.5, space_after=6)
                
            elif sec == "referees" and resume_data.get("referees"):
                p_sec = add_p("REFERENCES", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
                add_bottom_border(p_sec, border_color_hex)
                for ref in resume_data["referees"]:
                    add_p(f"• {ref}", size=9, space_after=2)

            elif sec == "cert_lang" and (certifications or languages):
                p_sec = add_p("CERTIFICATIONS & LANGUAGES", bold=True, size=12, color=accent_color, space_before=10, space_after=4)
                add_bottom_border(p_sec, border_color_hex)
                
                cert_items = []
                for cert in certifications:
                    c_name = safe_get(cert, "name")
                    c_issuer = safe_get(cert, "issuer")
                    issuer_str = f" ({c_issuer})" if c_issuer else ""
                    cert_items.append(f"{c_name}{issuer_str}")
                    
                lang_items = []
                for lang in languages:
                    l_name = safe_get(lang, "language")
                    prof = safe_get(lang, "proficiency")
                    prof_str = f" ({prof})" if prof else ""
                    lang_items.append(f"{l_name}{prof_str}")
                    
                if cert_items:
                    p_c = doc.add_paragraph()
                    p_c.paragraph_format.space_after = Pt(3)
                    r_lbl = p_c.add_run("Certifications: ")
                    r_lbl.bold = True
                    r_lbl.font.name = font_name
                    r_lbl.font.size = Pt(9.5)
                    r_val = p_c.add_run(", ".join(cert_items))
                    r_val.font.name = font_name
                    r_val.font.size = Pt(9.5)
                    
                if lang_items:
                    p_l = doc.add_paragraph()
                    p_l.paragraph_format.space_after = Pt(3)
                    r_lbl = p_l.add_run("Languages: ")
                    r_lbl.bold = True
                    r_lbl.font.name = font_name
                    r_lbl.font.size = Pt(9.5)
                    r_val = p_l.add_run(", ".join(lang_items))
                    r_val.font.name = font_name
                    r_val.font.size = Pt(9.5)

    # Save to a bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
